"""Script to create test files for multi-format parser agent"""

import json
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from PyPDF2 import PdfWriter
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


def create_pdf_file():
    """Create a sample PDF file"""
    print("Creating sample PDF file...")

    pdf_path = Path(__file__).parent / "sample.pdf"

    # Create PDF document
    doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
    story = []
    styles = getSampleStyleSheet()

    # Title
    title_style = ParagraphStyle(
        "CustomTitle", parent=styles["Heading1"], fontSize=24, textColor=colors.HexColor("#1a1a1a"), spaceAfter=30
    )
    story.append(Paragraph("Multi-Format Parser Agent Test Document", title_style))
    story.append(Spacer(1, 12))

    # Introduction
    story.append(Paragraph("Overview", styles["Heading2"]))
    intro_text = """
    This is a PDF document created to test the multi-format parser agent's PDF parsing capabilities.
    The parser should be able to extract text from multiple pages, read metadata, and handle
    various formatting elements including tables, lists, and styled text.
    """
    story.append(Paragraph(intro_text, styles["BodyText"]))
    story.append(Spacer(1, 20))

    # Features section
    story.append(Paragraph("Key Features Being Tested", styles["Heading2"]))
    features_list = [
        "Multi-page text extraction",
        "Metadata reading (title, author, subject)",
        "Table content extraction",
        "Formatted text handling",
        "Special characters and symbols: àáâãäå çèéê ñ ü",
    ]
    for feature in features_list:
        story.append(Paragraph(f"• {feature}", styles["BodyText"]))
    story.append(Spacer(1, 20))

    # Add a table
    story.append(Paragraph("Sample Data Table", styles["Heading2"]))
    data = [
        ["Format", "Extension", "Status"],
        ["PDF", ".pdf", "✓ Implemented"],
        ["DOCX", ".docx", "✓ Implemented"],
        ["CSV", ".csv", "✓ Implemented"],
        ["XLSX", ".xlsx", "✓ Implemented"],
        ["JSON", ".json", "✓ Implemented"],
        ["Markdown", ".md", "✓ Implemented"],
        ["Text", ".txt", "✓ Implemented"],
    ]

    table = Table(data)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 12),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 20))

    # Page break
    story.append(PageBreak())

    # Second page
    story.append(Paragraph("Page 2: Additional Content", styles["Heading2"]))
    page2_text = """
    This is the second page of the test PDF document. The multi-format parser should be able
    to extract text from all pages and combine them appropriately.
    """
    story.append(Paragraph(page2_text, styles["BodyText"]))
    story.append(Spacer(1, 20))

    # Technical specifications
    story.append(Paragraph("Technical Specifications", styles["Heading2"]))
    specs = """
    <b>Document Type:</b> Test PDF<br/>
    <b>Created:</b> 2025-11-17<br/>
    <b>Purpose:</b> Testing PDF parsing capabilities<br/>
    <b>Pages:</b> 2<br/>
    <b>Encoding:</b> UTF-8<br/>
    <b>Framework:</b> Agno AI<br/>
    """
    story.append(Paragraph(specs, styles["BodyText"]))
    story.append(Spacer(1, 20))

    # Conclusion
    story.append(Paragraph("Conclusion", styles["Heading2"]))
    conclusion = """
    This PDF document tests various PDF parsing features including multi-page extraction,
    table handling, formatted text, and metadata reading. The parser should successfully
    extract all content and present it in a structured, readable format.
    """
    story.append(Paragraph(conclusion, styles["BodyText"]))

    # Build PDF
    doc.build(story)
    print(f"✓ Created: {pdf_path}")


def create_docx_file():
    """Create a sample DOCX file"""
    print("Creating sample DOCX file...")

    docx_path = Path(__file__).parent / "sample.docx"

    # Create document
    doc = Document()

    # Title
    title = doc.add_heading("Multi-Format Parser Agent Test Document", 0)
    title.alignment = 1  # Center alignment

    # Introduction
    doc.add_heading("Overview", level=1)
    doc.add_paragraph(
        "This is a Microsoft Word document created to test the multi-format parser agent's "
        "DOCX parsing capabilities. The parser should be able to extract paragraphs, tables, "
        "and formatted text."
    )

    # Features
    doc.add_heading("Key Features Being Tested", level=1)
    features = doc.add_paragraph()
    features.add_run("The parser should handle:\n")
    doc.add_paragraph("Paragraph extraction", style="List Bullet")
    doc.add_paragraph("Table content extraction", style="List Bullet")
    doc.add_paragraph("Text formatting (bold, italic, underline)", style="List Bullet")
    doc.add_paragraph("Multiple heading levels", style="List Bullet")
    doc.add_paragraph("Special characters: àáâãäå çèéê ñ ü € £ ¥", style="List Bullet")

    # Add formatted text
    doc.add_heading("Formatted Text Example", level=1)
    p = doc.add_paragraph()
    p.add_run("This is normal text. ")
    p.add_run("This is bold text. ").bold = True
    p.add_run("This is italic text. ").italic = True
    p.add_run("This is underlined text. ").underline = True

    # Add a table
    doc.add_heading("Sample Data Table", level=1)
    table = doc.add_table(rows=8, cols=3)
    table.style = "Light Grid Accent 1"

    # Table headers
    headers = ["Format", "Extension", "Status"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    # Table data
    data = [
        ["PDF", ".pdf", "✓ Implemented"],
        ["DOCX", ".docx", "✓ Implemented"],
        ["CSV", ".csv", "✓ Implemented"],
        ["XLSX", ".xlsx", "✓ Implemented"],
        ["JSON", ".json", "✓ Implemented"],
        ["Markdown", ".md", "✓ Implemented"],
        ["Text", ".txt", "✓ Implemented"],
    ]

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_data

    # Technical specifications
    doc.add_heading("Technical Specifications", level=1)
    specs = doc.add_paragraph()
    specs.add_run("Document Type: ").bold = True
    specs.add_run("Test DOCX\n")
    specs.add_run("Created: ").bold = True
    specs.add_run("2025-11-17\n")
    specs.add_run("Purpose: ").bold = True
    specs.add_run("Testing DOCX parsing capabilities\n")
    specs.add_run("Framework: ").bold = True
    specs.add_run("Agno AI")

    # Conclusion
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "This DOCX document tests various Word document features including paragraphs, "
        "tables, formatted text, and multiple heading levels. The parser should successfully "
        "extract all content and present it in a structured format."
    )

    # Save document
    doc.save(str(docx_path))
    print(f"✓ Created: {docx_path}")


def create_xlsx_file():
    """Create a sample XLSX file"""
    print("Creating sample XLSX file...")

    xlsx_path = Path(__file__).parent / "sample.xlsx"

    # Create Excel writer
    with pd.ExcelWriter(str(xlsx_path), engine="openpyxl") as writer:
        # Sheet 1: Employee Data
        employee_data = {
            "Name": [
                "John Doe",
                "Jane Smith",
                "Mike Johnson",
                "Sarah Williams",
                "David Brown",
                "Emily Davis",
                "Chris Wilson",
                "Amanda Taylor",
            ],
            "Age": [32, 28, 45, 35, 29, 41, 33, 27],
            "Department": [
                "Engineering",
                "Marketing",
                "Engineering",
                "Sales",
                "Engineering",
                "Marketing",
                "Sales",
                "Engineering",
            ],
            "Salary": [95000, 78000, 125000, 88000, 82000, 92000, 95000, 76000],
            "Performance": [4.5, 4.8, 4.2, 4.6, 4.3, 4.7, 4.4, 4.1],
        }
        df_employees = pd.DataFrame(employee_data)
        df_employees.to_excel(writer, sheet_name="Employees", index=False)

        # Sheet 2: Department Summary
        dept_summary = {
            "Department": ["Engineering", "Marketing", "Sales"],
            "Headcount": [4, 2, 2],
            "Avg_Salary": [94500, 85000, 91500],
            "Avg_Performance": [4.28, 4.75, 4.50],
            "Budget": [378000, 170000, 183000],
        }
        df_dept = pd.DataFrame(dept_summary)
        df_dept.to_excel(writer, sheet_name="Department_Summary", index=False)

        # Sheet 3: Quarterly Metrics
        quarterly_data = {
            "Quarter": ["Q1 2024", "Q2 2024", "Q3 2024", "Q4 2024"],
            "Revenue": [1250000, 1380000, 1420000, 1550000],
            "Expenses": [890000, 920000, 950000, 980000],
            "Profit": [360000, 460000, 470000, 570000],
            "Growth_%": [5.2, 10.4, 2.9, 9.0],
        }
        df_quarterly = pd.DataFrame(quarterly_data)
        df_quarterly.to_excel(writer, sheet_name="Quarterly_Metrics", index=False)

    print(f"✓ Created: {xlsx_path}")


def main():
    """Create all test files"""
    print("=" * 60)
    print("Creating Test Files for Multi-Format Parser Agent")
    print("=" * 60)

    try:
        create_pdf_file()
        create_docx_file()
        create_xlsx_file()

        print("=" * 60)
        print("✓ All test files created successfully!")
        print("=" * 60)
        print("\nCreated files:")
        test_dir = Path(__file__).parent
        for file in sorted(test_dir.glob("sample.*")):
            if file.suffix in [".pdf", ".docx", ".xlsx", ".csv", ".json", ".md", ".txt"]:
                size = file.stat().st_size
                print(f"  • {file.name} ({size:,} bytes)")

    except Exception as e:
        print(f"\n❌ Error creating test files: {e}")
        import traceback

        traceback.print_exc()


if __name__ == "__main__":
    main()
