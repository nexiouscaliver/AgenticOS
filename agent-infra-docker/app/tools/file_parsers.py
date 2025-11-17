"""Multi-format file parser tools for parsing PDF, DOCX, CSV, XLSX, JSON, MD, and TXT files"""

import json
import os
from pathlib import Path
from textwrap import dedent
from typing import Optional

import chardet
import pandas as pd
from agno.tools import Toolkit
from agno.utils.log import logger
from docx import Document
from markdown import markdown
from PyPDF2 import PdfReader


class MultiFormatFileParser(Toolkit):
    """Toolkit for parsing multiple file formats including PDF, DOCX, CSV, XLSX, JSON, MD, and TXT"""

    def __init__(
        self,
        max_file_size_mb: int = 50,
        csv_max_rows: int = 10000,
        xlsx_max_rows: int = 10000,
    ):
        """
        Initialize the multi-format file parser

        Args:
            max_file_size_mb: Maximum file size in MB to process
            csv_max_rows: Maximum rows to read from CSV files
            xlsx_max_rows: Maximum rows to read from XLSX files
        """
        super().__init__(name="multi_format_file_parser")

        self.max_file_size_bytes = max_file_size_mb * 1024 * 1024
        self.csv_max_rows = csv_max_rows
        self.xlsx_max_rows = xlsx_max_rows

        self.register(self.parse_file)
        self.register(self.list_supported_formats)
        self.register(self.get_file_info)

    def _check_file_size(self, file_path: str) -> tuple[bool, str]:
        """Check if file size is within limits"""
        try:
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size_bytes:
                size_mb = file_size / (1024 * 1024)
                max_mb = self.max_file_size_bytes / (1024 * 1024)
                return False, f"File size ({size_mb:.2f}MB) exceeds maximum allowed size ({max_mb}MB)"
            return True, ""
        except Exception as e:
            return False, f"Error checking file size: {str(e)}"

    def _detect_encoding(self, file_path: str) -> str:
        """Detect file encoding using chardet"""
        try:
            with open(file_path, "rb") as file:
                raw_data = file.read(10000)  # Read first 10KB
                result = chardet.detect(raw_data)
                return result["encoding"] or "utf-8"
        except Exception:
            return "utf-8"

    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF file and extract text"""
        try:
            reader = PdfReader(file_path)
            text_content = []

            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(f"--- Page {page_num} ---\n{page_text}")

            if not text_content:
                return "⚠️ PDF file contains no extractable text (might be image-based)"

            metadata = reader.metadata
            result = f"📄 PDF Document ({len(reader.pages)} pages)\n\n"

            if metadata:
                result += "Metadata:\n"
                if metadata.title:
                    result += f"  Title: {metadata.title}\n"
                if metadata.author:
                    result += f"  Author: {metadata.author}\n"
                if metadata.subject:
                    result += f"  Subject: {metadata.subject}\n"
                result += "\n"

            result += "Content:\n" + "\n\n".join(text_content)
            return result

        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            return f"❌ Error parsing PDF file: {str(e)}"

    def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX file and extract text"""
        try:
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            if not paragraphs:
                return "⚠️ DOCX file contains no text content"

            # Extract tables if any
            tables_content = []
            for table_num, table in enumerate(doc.tables, 1):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                if table_data:
                    tables_content.append(f"Table {table_num}:\n" + "\n".join(table_data))

            result = f"📝 DOCX Document ({len(paragraphs)} paragraphs"
            if tables_content:
                result += f", {len(tables_content)} tables"
            result += ")\n\n"

            result += "Content:\n" + "\n\n".join(paragraphs)

            if tables_content:
                result += "\n\n--- Tables ---\n\n" + "\n\n".join(tables_content)

            return result

        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            return f"❌ Error parsing DOCX file: {str(e)}"

    def _parse_csv(self, file_path: str) -> str:
        """Parse CSV file and convert to structured text"""
        try:
            encoding = self._detect_encoding(file_path)
            df = pd.read_csv(file_path, encoding=encoding, nrows=self.csv_max_rows)

            rows, cols = df.shape
            result = f"📊 CSV File ({rows} rows × {cols} columns)\n\n"

            result += "Columns: " + ", ".join(df.columns.tolist()) + "\n\n"

            # Summary statistics for numeric columns
            numeric_cols = df.select_dtypes(include=["number"]).columns
            if len(numeric_cols) > 0:
                result += "Numeric Summary:\n"
                summary = df[numeric_cols].describe()
                result += summary.to_string() + "\n\n"

            # Show first few rows
            result += f"First {min(10, rows)} rows:\n"
            result += df.head(10).to_string(index=False) + "\n"

            if rows == self.csv_max_rows:
                result += f"\n⚠️ Only first {self.csv_max_rows} rows loaded (file may contain more)"

            return result

        except Exception as e:
            logger.error(f"Error parsing CSV: {e}")
            return f"❌ Error parsing CSV file: {str(e)}"

    def _parse_xlsx(self, file_path: str) -> str:
        """Parse XLSX file and convert to structured text"""
        try:
            # Get all sheet names
            xl_file = pd.ExcelFile(file_path)
            sheet_names = xl_file.sheet_names

            result = f"📊 Excel File ({len(sheet_names)} sheet(s))\n\n"
            result += "Sheets: " + ", ".join(sheet_names) + "\n\n"

            # Parse each sheet
            for sheet_name in sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name, nrows=self.xlsx_max_rows)
                rows, cols = df.shape

                result += f"--- Sheet: {sheet_name} ({rows} rows × {cols} columns) ---\n\n"

                if cols > 0:
                    result += "Columns: " + ", ".join(df.columns.astype(str).tolist()) + "\n\n"

                    # Summary statistics for numeric columns
                    numeric_cols = df.select_dtypes(include=["number"]).columns
                    if len(numeric_cols) > 0:
                        result += "Numeric Summary:\n"
                        summary = df[numeric_cols].describe()
                        result += summary.to_string() + "\n\n"

                    # Show first few rows
                    result += f"First {min(10, rows)} rows:\n"
                    result += df.head(10).to_string(index=False) + "\n\n"

                    if rows == self.xlsx_max_rows:
                        result += f"⚠️ Only first {self.xlsx_max_rows} rows loaded from this sheet\n\n"

            return result

        except Exception as e:
            logger.error(f"Error parsing XLSX: {e}")
            return f"❌ Error parsing XLSX file: {str(e)}"

    def _parse_json(self, file_path: str) -> str:
        """Parse JSON file and format content"""
        try:
            encoding = self._detect_encoding(file_path)
            with open(file_path, "r", encoding=encoding) as f:
                data = json.load(f)

            # Format JSON with indentation
            formatted_json = json.dumps(data, indent=2, ensure_ascii=False)

            # Get basic info
            data_type = type(data).__name__
            if isinstance(data, dict):
                size_info = f"{len(data)} keys"
            elif isinstance(data, list):
                size_info = f"{len(data)} items"
            else:
                size_info = "1 value"

            result = f"🔧 JSON File ({data_type}, {size_info})\n\n"
            result += "Content:\n" + formatted_json

            return result

        except json.JSONDecodeError as e:
            logger.error(f"Error parsing JSON: {e}")
            return f"❌ Invalid JSON file: {str(e)}"
        except Exception as e:
            logger.error(f"Error parsing JSON: {e}")
            return f"❌ Error parsing JSON file: {str(e)}"

    def _parse_markdown(self, file_path: str) -> str:
        """Parse Markdown file"""
        try:
            encoding = self._detect_encoding(file_path)
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()

            if not content.strip():
                return "⚠️ Markdown file is empty"

            # Count sections (lines starting with #)
            sections = [line for line in content.split("\n") if line.strip().startswith("#")]

            result = f"📖 Markdown File ({len(sections)} headings)\n\n"
            result += "Raw Content:\n" + content

            return result

        except Exception as e:
            logger.error(f"Error parsing Markdown: {e}")
            return f"❌ Error parsing Markdown file: {str(e)}"

    def _parse_text(self, file_path: str) -> str:
        """Parse plain text file"""
        try:
            encoding = self._detect_encoding(file_path)
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()

            if not content.strip():
                return "⚠️ Text file is empty"

            lines = content.split("\n")
            word_count = len(content.split())
            char_count = len(content)

            result = f"📄 Text File ({len(lines)} lines, {word_count} words, {char_count} characters)\n\n"
            result += "Content:\n" + content

            return result

        except Exception as e:
            logger.error(f"Error parsing text file: {e}")
            return f"❌ Error parsing text file: {str(e)}"

    def parse_file(self, file_path: str) -> str:
        """
        Parse a file and extract its content based on file type

        Supports: PDF, DOCX, CSV, XLSX, JSON, MD, TXT

        Args:
            file_path: Path to the file to parse (can be relative or absolute)

        Returns:
            Extracted and formatted content from the file
        """
        try:
            # Resolve path
            path = Path(file_path).resolve()

            if not path.exists():
                return f"❌ File not found: {file_path}"

            if not path.is_file():
                return f"❌ Path is not a file: {file_path}"

            # Check file size
            size_ok, size_msg = self._check_file_size(str(path))
            if not size_ok:
                return f"❌ {size_msg}"

            # Get file extension
            extension = path.suffix.lower()

            # Route to appropriate parser
            parsers = {
                ".pdf": self._parse_pdf,
                ".docx": self._parse_docx,
                ".doc": self._parse_docx,
                ".csv": self._parse_csv,
                ".xlsx": self._parse_xlsx,
                ".xls": self._parse_xlsx,
                ".json": self._parse_json,
                ".md": self._parse_markdown,
                ".markdown": self._parse_markdown,
                ".txt": self._parse_text,
            }

            parser = parsers.get(extension)
            if not parser:
                return dedent(
                    f"""
                    ❌ Unsupported file format: {extension}

                    Supported formats: {", ".join(parsers.keys())}

                    Use list_supported_formats() for more details.
                    """
                ).strip()

            logger.info(f"Parsing {extension} file: {path}")
            result = parser(str(path))

            return result

        except Exception as e:
            logger.error(f"Error in parse_file: {e}")
            return f"❌ Unexpected error parsing file: {str(e)}"

    def list_supported_formats(self) -> str:
        """
        List all supported file formats and their descriptions

        Returns:
            Information about supported file formats
        """
        return dedent(
            """
            📋 Supported File Formats:

            📄 PDF (.pdf)
               - Extracts text content from all pages
               - Includes document metadata (title, author, etc.)
               - Multi-page support

            📝 Microsoft Word (.docx, .doc)
               - Extracts paragraphs and text content
               - Includes tables if present
               - Preserves document structure

            📊 CSV (.csv)
               - Parses tabular data
               - Shows column names and data types
               - Includes summary statistics for numeric columns
               - Preview of first 10 rows

            📊 Excel (.xlsx, .xls)
               - Supports multiple sheets
               - Parses tabular data per sheet
               - Shows column names and summary statistics
               - Preview of first 10 rows per sheet

            🔧 JSON (.json)
               - Parses structured JSON data
               - Pretty-printed formatting
               - Handles nested objects and arrays

            📖 Markdown (.md, .markdown)
               - Extracts markdown content
               - Preserves formatting syntax
               - Counts headings and sections

            📄 Plain Text (.txt)
               - Reads text content with encoding detection
               - Shows line, word, and character counts
               - Supports various text encodings

            Configuration:
            - Max file size: 50MB (configurable)
            - CSV/XLSX max rows: 10,000 per sheet (configurable)
            - Automatic encoding detection for text files
            """
        ).strip()

    def get_file_info(self, file_path: str) -> str:
        """
        Get basic information about a file without parsing it

        Args:
            file_path: Path to the file

        Returns:
            File information (name, size, type, etc.)
        """
        try:
            path = Path(file_path).resolve()

            if not path.exists():
                return f"❌ File not found: {file_path}"

            if not path.is_file():
                return f"❌ Path is not a file: {file_path}"

            size_bytes = path.stat().st_size
            size_mb = size_bytes / (1024 * 1024)
            extension = path.suffix.lower()

            result = dedent(
                f"""
                📄 File Information:

                Name: {path.name}
                Path: {path}
                Size: {size_bytes:,} bytes ({size_mb:.2f} MB)
                Extension: {extension}
                Type: {self._get_format_name(extension)}
                """
            ).strip()

            return result

        except Exception as e:
            logger.error(f"Error getting file info: {e}")
            return f"❌ Error getting file information: {str(e)}"

    def _get_format_name(self, extension: str) -> str:
        """Get human-readable format name"""
        formats = {
            ".pdf": "PDF Document",
            ".docx": "Microsoft Word Document",
            ".doc": "Microsoft Word Document (Legacy)",
            ".csv": "CSV (Comma-Separated Values)",
            ".xlsx": "Microsoft Excel Spreadsheet",
            ".xls": "Microsoft Excel Spreadsheet (Legacy)",
            ".json": "JSON Data",
            ".md": "Markdown Document",
            ".markdown": "Markdown Document",
            ".txt": "Plain Text",
        }
        return formats.get(extension, "Unknown")
